# 读取docx中的文本代码示例
import re
import xlwt
import docx

# 获取文档对象
file = docx.Document("1.docx")
print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段
book = xlwt.Workbook(encoding='utf-8', style_compression=0)
sheet = book.add_sheet('word', cell_overwrite_ok=True)

# 输出每一段的内容
num = 0
for para in file.paragraphs:
    # print(para.text)
    zn = str(para.text)
    words = re.findall(r'(.+)（', zn)
    words = ' '.join(words)
    if words != "":
        try:
            print(words)
            sheet.write(num, 0, words)
            num = num + 1
        except:
            pass
    else:
        pass
print(num)

#num = 0
#for para in file.paragraphs:
    # print(para.text)
   # zn = str(para.text)
    #words = re.match(r'([a-zA-Z\s])+(.+)', zn)
   # try:
       # print(words.group(2))
        #sheet.write(num, 1, words.group(2))
        #num = num + 1
   # except:
       # pass

# 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
# print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)


savepath = r'1.xls'
book.save(savepath)
